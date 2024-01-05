import openai
import subprocess
import datetime
import json
from docx import Document
import sqlite3
import hashlib
import sys

def createfilename(text):
    symbols = (u"абвгдеёжзийклмнопрстуфхцчшщъыьэюяАБВГДЕЁЖЗИЙКЛМНОПРСТУФХЦЧШЩЪЫЬЭЮЯ",
               u"abvgdeejzijklmnoprstufhzcss_y_euaABVGDEEJZIJKLMNOPRSTUFHZCSS_Y_EUA")
    tr = {ord(a): ord(b) for a, b in zip(*symbols)}
    return text.translate(tr).replace(" ", "_").replace("-", "_").replace(".", "_").replace("/", "_").replace(":", "_")

def save_as_docx(text, output_file):
    doc = Document()
    doc.add_paragraph(text)
    doc.save(output_file)


def beep(frequency=1000, duration=500):
    cmd = f'play -n synth {duration / 1000} sin {frequency} vol 0.5'
    subprocess.run(cmd, shell=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)

def docx_to_txt(docx_path):
    document = Document(docx_path)
    text = ''
    for paragraph in document.paragraphs:
        text += paragraph.text + '\n'

    return text


def docx_replace(old_file, new_file, replacements):
    doc = Document(old_file)
    for paragraph in doc.paragraphs:
        if replacements:
            for key in replacements.keys():
                if key in paragraph.text:
                    paragraph.text = paragraph.text.replace(key, replacements[key])
    doc.save(new_file)


def prep_rep(JD):
    global run_id
    replacements = {}
    reps = ["__rep1", "__rep2"]
    global totalrate
    for rep in reps:

        # repsa = []
        # with open(rep + '.txt', 'r') as ffile:
        #     for line in ffile:
        #         repsa[line] = 0
        reprate = 0
        repsa = set()
        with open(rep + '.txt', 'r') as f:
            for id, line in enumerate(f, start=1):
                # name = line.strip()  # Remove newline characters
                repsa.add((id, line, 0))

        # print('Repsa vvvvvv')
        # for item in repsa:
        #     print(item)
        # print('Repsa ^^^^^^')
        repsa_str = json.dumps(list(repsa))
        prompt = f"""
Imagine you are an AI-based hiring manager. Your task is to evaluate the relevance of various bullet points of a candidate's bullet points {repsa_str} against a provided job description {JD}. The bullet points is in the form of string representations of list, having fileds (id, name, score).
Score each bullet point on float a scale from 0 to 10 with with details to 0.1 based on how closely it matches with the job description. 
Return this data in the form of a JSON object where each key corresponds to the id in double quotes, and the value is the score you assigned.
"""

        resp_txt = fast_response(prompt)
        try:
            data_dict = json.loads(resp_txt)
        except json.JSONDecodeError as e:
            print("Failed to decode JSON: ", str(e))
            print("JSON:", resp_txt)
            beep(700, 450)
            beep(700, 450)
            beep(700, 450)
            exit(1)

        totalpoint = ''
        tuple_dict = {item[0]: item for item in repsa}
        ic = 0
        for key, value in sorted(data_dict.items(), key=lambda x: float(x[1]), reverse=True):
            pn = tuple_dict[int(key)][1]
            id_var = int(hashlib.md5(pn.encode()).hexdigest()[:15], 16)

            sqlcc.execute("SELECT id FROM bup WHERE id = ?", (id_var,))
            existing_id = sqlcc.fetchone()
            if existing_id is None:
                values = (id_var, pn, rep, run_id)
                sqlcc.execute("INSERT INTO bup (id, nm, em, run_id) VALUES (?, ?, ?, ?)", values)
                bup_id = sqlcc.lastrowid
                is_new="Y"
            else:
                bup_id = int(existing_id[0])
                is_new='N'


            if (rep == "__rep1"): max_ic=5
            elif (rep == "__rep2"): max_ic=4
            else: max_ic=3

            ic += 1
            if ic <= max_ic:
                totalpoint = totalpoint + '- ' + pn
                #print('-A-', value, '<', id_var, '>', key, is_new,' ', pn, end='')
                decision = 1
                reprate+=value
            else:
                #print('-D-', value, '<', id_var, '>', key, is_new, ' ', pn, end='')
                decision = 0

            values = (value, decision, run_id, bup_id)
            sqlcc.execute("INSERT INTO sco (sc, de, run_id, bup_id) VALUES (?, ?, ?, ?)", values)


        print('^---',reprate,'----=|', rep, '                   |=-----------------^')
        replacements[rep] = totalpoint
        totalrate+=reprate
    print('^--',totalrate,'----=| Total                     |=-----------------^')
    return replacements


def fast_response(myprompt):
    beep(300, 250)     #"gpt-3.5-turbo"
    chat_completion = openai.ChatCompletion.create(model="gpt-3.5-turbo",
                                                   messages=[
                                                       {"role": "user", "content": myprompt}
                                                   ],
                                                   temperature=1
                                                   )
    return chat_completion.choices[0].message.content

#____________________________________________________________________________________________________________
totalrate=0
with open('openai.key','r') as key_file:
    openai.api_key=key_file.read().strip()

conn = sqlite3.connect('log.db')
sqlcc = conn.cursor()

with open('JD.txt', 'r') as file:
    JDdoc = file.read()

JD = f"<Job description START>{JDdoc}<Job description END>"
# print(JD)

with open('CV.txt', 'r') as file:
    CVdoc = file.read()
CV = f"<Resume START>{CVdoc}<Resume END>"
# print(CV)

prompt = f"""
Here is Job descirption {JD} Can you please look into it and return to me only JSON object with keys "company_name" "position_name" keys containgin company name and position name repsectively.
"""
resp_txt = fast_response(prompt)
try:
    data_dict = json.loads(resp_txt)
except json.JSONDecodeError as e:
    print("Failed to decode JSON: ", str(e))
    print("JSON:", resp_txt)
    beep(700, 450)
    beep(700, 450)
    beep(700, 450)
    exit(1)

for key, value in data_dict.items():
    if key == "company_name": company_name=value;
    if key == "position_name": position_name = value;
print(company_name,' ',position_name)
print('________________________________________________________')

sqlcc.execute("SELECT * FROM run WHERE com LIKE ?", ('%' + company_name + '%',))
rows = sqlcc.fetchall()
requires_confirm_toproceed=0
for row in rows:
    if row[5] == position_name:
        beep(500, 350)
        beep(500, 350)
        utc_time = datetime.datetime.strptime(row[1], '%Y-%m-%d %H:%M:%S')
        pst_time = utc_time - datetime.timedelta(hours=8)
        print("\033[91m", pst_time, row[4], row[5], "\033[0m")
        requires_confirm_toproceed = 1
    else:
        print('',row[1],row[4],row[5])

if requires_confirm_toproceed:
    input("Press Enter to continue, or Ctrl+C to abort!")

values = (JDdoc, CVdoc, company_name, position_name)
sqlcc.execute("INSERT INTO run (dt, jd, cv, com, pos) VALUES (datetime('now'), ?, ?, ?, ?)", values)
run_id = sqlcc.lastrowid

replacements = prep_rep(JD)
# print('----V-----')
# print(replacements)
# print('----^-----')

current_time = datetime.datetime.now()
mfilename = createfilename("CV_"+str(int(totalrate))+"_"+company_name+"_"+current_time.strftime("_%H_%M%S"))
old_file = "__cv.docx"
docx_replace(old_file, mfilename + ".docx", replacements)

docx_file_path = mfilename + ".docx"
txt_content = docx_to_txt(docx_file_path)


with open('CV.txt', 'w') as txt_file:
    txt_file.write(txt_content)
# with open(filename+".txt", "w") as file:
#    file.write(result)
# save_as_docx(result, filename + ".docx")
# pyperclip.copy(result)
sqlcc.execute("UPDATE run SET cv = ? WHERE id = ?", (txt_content, run_id))

conn.commit()
conn.close()

beep(400, 150)
beep(500, 150)
beep(300, 350)
print('<-----------=| All done: ' + mfilename + '.docx |=----------->')
with open('run_id.txt', 'w') as txt_file:
    txt_file.write(str(run_id))
print('-= Resume saved =-')